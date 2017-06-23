from docx import Document
import docx
import re
import os
from tkinter import *
from tkinter import filedialog
from tkinter import messagebox

# regolar expression usata
regex = r"(?s)((if\s).*?(then).+?(?=else|if)|(else).+?(\n))"


def gui():

    def button_go_callback():
        """ what to do when the "Go" button is pressed """
        filedocx = entry.get()
        if filedocx.rsplit(".")[-1] != "docx":
            statusText.set("Filename must end in `.docx'")
            message.configure(fg="red")
            return
        else:

            # trasformo il testo da docx a testo trattabile e applico la regex
            test_str = getText(filedocx)
            lenght = len(test_str)
            print('Il documento analizzato contiene %d caratteri' % lenght)
            matches = re.finditer(regex, test_str, re.IGNORECASE | re.VERBOSE | re.MULTILINE)

            # costruttore docx
            doc = Document()
            doc.add_heading('Result', 0)
            size = 0

            # ciclo for che tramite regex mi individua il costrutto if then else
            for matchNum, match in enumerate(matches):
                matchNum = matchNum + 1

                m = ("\n ********************** Match {matchNum} was found at {start}-{end} **********************: "
                     "\n {match}".format(matchNum=matchNum, start=match.start(), end=match.end(), match=match.group()))
                # print(m.encode("utf-8"))

                # calcolo dell'estrazione parziale e aggiornamento dell'estrazione totale
                lengroup = (match.end() - match.start())
                print("\n Match lungo %d caratteri" % lengroup)
                size = size + lengroup
                print("\n In totale sono stati estratti %d caratteri " % size)
                percent = size / lenght * 100
                print("\n Percentuale di documento analizzata %s per cento" % round(percent, 2))

                # scrittura su file docx di output
                doc.add_paragraph(m)
                # font = p.add_run(match.group()).font

            # file di output
            doc.add_heading('Stats', 0)
            doc.add_paragraph('Il documento analizzato contiene %d caratteri' % (len(test_str)))
            doc.add_paragraph('In totale sono stati estratti %d caratteri' % size)
            doc.add_paragraph('Percentuale del documento trattata: %s ' % round(percent, 2))
            output_file = entry2.get()
            doc.save(output_file)
            messagebox.showinfo(message='Analisi completata con successo e salvataggio del file di output effettuato')

    def button_output_callback():
        #output_file = filedialog.asksaveasfilename(defaultextension=" .docx ", filetypes=[('word files', ' .docx ')])
        suffix = '_ITE_Out.docx'
        base_filename = entry.get()
        output_file = os.path.join(base_filename + suffix)
        entry2.delete(0, END)
        entry2.insert(0, output_file)
        return output_file

    def button_browse_callback():
        """ What to do when the Browse button is pressed """
        filedocx = filedialog.askopenfilename()
        #suffix_in = '.docx'
        #base_in_filename = filedialog.askopenfilename()
        #filedocx = os.path.join(base_in_filename + suffix_in)
        entry.delete(0, END)
        entry.insert(0, filedocx)
        return filedocx

    root = Tk(className="ITExtract")
    root.geometry("640x640")
    frame = Frame(root)
    fname = Canvas(height=300, width=500)
    fname.pack(side=TOP)
    logo = PhotoImage(file="ITExtractlogo.jpg")
    fname.create_image(250, 200, anchor=CENTER, image=logo)
    frame.pack()

    statusText = StringVar(root)
    statusText.set("1) Press Input file to choose input docx file, "
                   "\n 2) press Output file to choose output file name,"
                   "\n  3) press ITExtract button")

    label = Label(root,)
    label.pack()
    entry = Entry(root, width=80)

    label2 = Label(root,)
    label2.pack()
    entry2 = Entry(root, width=80)

    separator = Frame(root, height=2, bd=1, relief=SUNKEN)
    separator.pack(fill=X, padx=5, pady=5)

    button_browse = Button(root,
                           text="Input file Browse",
                           command=button_browse_callback)
    button_out = Button(root,
                        text="Output file Browse",
                        command=button_output_callback)
    button_browse.pack()
    entry.pack()
    button_out.pack()
    entry2.pack()
    separator = Frame(root, height=2, bd=1, relief=SUNKEN)
    separator.pack(fill=X, padx=5, pady=5)
    button_go = Button(root,
                       text="ITExtract",
                       command=button_go_callback,
                       anchor=NE)
    button_go.pack()

    message = Label(root, textvariable=statusText)
    message.pack()

    mainloop()

def getText(filename):
    doc = docx.Document(filename)
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return '\n'.join(fulltext)

if __name__ == "__main__":
    """ Run as a stand-alone script """

    gui()
